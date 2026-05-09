"""
Word (.docx) 보고서 생성기.

구성:
  표지 → 분석 개요 테이블 → 데이터 시각화(차트) → 10개 섹션 → 출처 목록
이메일 첨부 후 사용자가 직접 PPT로 변환 가능한 형태로 작성.
"""
import io
import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn as docx_qn
from docx.shared import Inches, Pt, RGBColor

# ── 색상 ─────────────────────────────────────────────────────────────────────
C_NAVY  = RGBColor(0x02, 0x3E, 0x8A)
C_CYAN  = RGBColor(0x00, 0xB4, 0xD8)
C_DARK  = RGBColor(0x0D, 0x1B, 0x2A)
C_GRAY  = RGBColor(0x55, 0x66, 0x77)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_TEXT  = RGBColor(0x1A, 0x1A, 0x2E)

# 섹션별 키 컬러 (Navy → Cyan 계열)
_SEC_COLORS = [
    RGBColor(0x02, 0x3E, 0x8A), RGBColor(0x00, 0x77, 0xB6),
    RGBColor(0x00, 0x96, 0xC7), RGBColor(0x00, 0xB4, 0xD8),
    RGBColor(0x48, 0xCA, 0xE4), RGBColor(0x02, 0x3E, 0x8A),
    RGBColor(0x00, 0x77, 0xB6), RGBColor(0x00, 0x96, 0xC7),
    RGBColor(0x00, 0xB4, 0xD8), RGBColor(0x48, 0xCA, 0xE4),
]


# ══════════════════════════════════════════════════════════════════════════════
# 내부 유틸
# ══════════════════════════════════════════════════════════════════════════════

def _clean(text: str) -> str:
    """마크다운 제거 → 순수 텍스트."""
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)   # 링크
    text = re.sub(r"\*{1,2}([^*]+)\*{1,2}", r"\1", text)    # 볼드/이탤릭
    text = re.sub(r"^#{1,4}\s*", "", text, flags=re.MULTILINE)  # 헤딩 #
    return text.strip()


def _set_cell_bg(cell, hex_color: str) -> None:
    """테이블 셀 배경색."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(docx_qn("w:val"),   "clear")
    shd.set(docx_qn("w:color"), "auto")
    shd.set(docx_qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def _cell_text(cell, text: str, bold: bool = False,
               color: RGBColor = C_TEXT, size_pt: float = 10) -> None:
    """테이블 셀 텍스트 설정."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.bold = bold
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color


def _heading(doc: Document, text: str, level: int = 1,
             color: RGBColor = C_NAVY) -> None:
    p = doc.add_heading(text, level=level)
    if p.runs:
        run = p.runs[0]
    else:
        run = p.add_run(text)
    run.font.color.rgb = color
    run.font.bold = True


def _caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(8.5)
        run.font.italic = True
        run.font.color.rgb = C_GRAY
    p.paragraph_format.space_after = Pt(10)


def _embed_image(doc: Document, img_bytes: bytes,
                 width_in: float = 5.5, caption: str = "") -> None:
    """PNG bytes를 문서에 삽입."""
    if not img_bytes:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(io.BytesIO(img_bytes), width=Inches(width_in))
    if caption:
        _caption(doc, caption)


def _section_bullets(doc: Document, content: str,
                     accent: RGBColor = C_NAVY) -> None:
    """섹션 내용 → 불릿 단락 변환."""
    for raw in content.split("\n"):
        line = _clean(raw.strip().lstrip("-•·▪▸").strip())
        if not line:
            continue
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.color.rgb = C_TEXT
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.space_before = Pt(2)


# ══════════════════════════════════════════════════════════════════════════════
# 메인 함수
# ══════════════════════════════════════════════════════════════════════════════

def generate_word(analysis: dict, output_dir: str = "/tmp") -> str:
    """
    분석 결과 dict를 Word 문서로 저장.
    Returns: 저장된 .docx 파일 경로
    """
    from .chart_generator import (
        make_keyword_chart_bytes,
        make_source_chart_bytes,
        make_workload_chart_bytes,
    )

    # ── 메타 추출 ─────────────────────────────────────────────
    date_str    = analysis.get("date", datetime.now(timezone.utc).strftime("%Y-%m-%d"))
    report_type = analysis.get("type", "daily").upper()
    attribution = analysis.get("model_attribution", analysis.get("provider", "-"))
    art_count   = analysis.get("article_count", 0)
    keywords    = analysis.get("keywords_found", [])
    sources     = analysis.get("sources", [])
    source_urls = analysis.get("source_urls", [])
    sections    = analysis.get("sections", {})
    week_label  = analysis.get("week_label", "")
    subtitle    = week_label or date_str
    now_utc     = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    doc = Document()

    # 페이지 여백
    sec = doc.sections[0]
    sec.top_margin    = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin   = Inches(1.1)
    sec.right_margin  = Inches(1.1)

    # ══════════════════════════════════════════════════════════
    # 1. 표지
    # ══════════════════════════════════════════════════════════
    for _ in range(3):
        doc.add_paragraph()

    # 메인 타이틀
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("AI/Semiconductor Daily News")
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = C_NAVY

    p_title2 = doc.add_paragraph()
    p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_title2.add_run("Automotive/Humanoid and Storage Intelligence")
    r2.font.size = Pt(14)
    r2.font.bold = False
    r2.font.color.rgb = C_CYAN

    doc.add_paragraph()

    # 리포트 타입 + 날짜
    p_type = doc.add_paragraph()
    p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p_type.add_run(f"[{report_type}]  {subtitle}")
    r3.font.size = Pt(16)
    r3.font.color.rgb = C_CYAN

    doc.add_paragraph()
    doc.add_paragraph()

    # 메타 정보
    meta_lines = [
        f"수집 기사: {art_count}건",
        f"핵심 키워드: {', '.join(keywords[:6])}",
        f"AI 기여 모델: {attribution}",
        f"생성 일시: {now_utc}",
    ]
    for line in meta_lines:
        p_m = doc.add_paragraph()
        p_m.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_m = p_m.add_run(line)
        r_m.font.size = Pt(10.5)
        r_m.font.color.rgb = C_GRAY

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 2. 분석 개요 테이블
    # ══════════════════════════════════════════════════════════
    _heading(doc, "분석 개요 (Overview)", level=1)

    rows_data = [
        ("항목",          "내용"),           # 헤더
        ("📅 분석 일자",  subtitle),
        ("📰 수집 기사",  f"{art_count}건"),
        ("🏷 핵심 키워드", ", ".join(keywords[:8])),
        ("📡 데이터 출처", ", ".join(sources[:6])),
        ("🤖 AI 기여",    attribution),
    ]

    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.style = "Table Grid"

    for i, (label, value) in enumerate(rows_data):
        row = tbl.rows[i]
        if i == 0:                    # 헤더 행
            _set_cell_bg(row.cells[0], "023E8A")
            _set_cell_bg(row.cells[1], "023E8A")
            _cell_text(row.cells[0], label, bold=True, color=C_WHITE, size_pt=10)
            _cell_text(row.cells[1], value, bold=True, color=C_WHITE, size_pt=10)
        else:
            _set_cell_bg(row.cells[0], "EBF5FB")
            _cell_text(row.cells[0], label, bold=True, color=C_NAVY, size_pt=10)
            _cell_text(row.cells[1], value, color=C_TEXT, size_pt=10)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # 3. 데이터 시각화 (차트)
    # ══════════════════════════════════════════════════════════
    _heading(doc, "데이터 시각화", level=1)

    # 키워드 차트
    kw_bytes = make_keyword_chart_bytes(keywords)
    if kw_bytes:
        _embed_image(doc, kw_bytes, width_in=5.5,
                     caption="[그림 1] 핵심 키워드 우선순위 (분석 빈도 기반)")

    # 출처 분포 차트
    src_bytes = make_source_chart_bytes(sources)
    if src_bytes:
        _embed_image(doc, src_bytes, width_in=4.5,
                     caption="[그림 2] 데이터 출처 분포")

    # 스토리지 워크로드 차트 (8번 섹션 관련)
    wl_bytes = make_workload_chart_bytes()
    if wl_bytes:
        _embed_image(doc, wl_bytes, width_in=4.5,
                     caption="[그림 3] 차량용 스토리지 Read/Write 비율 (ADAS 기준)")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 4. 10개 분석 섹션
    # ══════════════════════════════════════════════════════════
    _heading(doc, "분석 결과 (Analysis)", level=1)
    doc.add_paragraph()

    for idx, (sec_title, sec_content) in enumerate(sections.items()):
        accent = _SEC_COLORS[idx % len(_SEC_COLORS)]

        # 섹션 제목
        p_sec = doc.add_heading(sec_title, level=2)
        if p_sec.runs:
            run_sec = p_sec.runs[0]
        else:
            run_sec = p_sec.add_run(sec_title)
        run_sec.font.color.rgb = accent
        run_sec.font.bold = True
        p_sec.paragraph_format.space_before = Pt(14)
        p_sec.paragraph_format.space_after  = Pt(6)

        # 섹션 내용 → 불릿
        _section_bullets(doc, sec_content, accent=accent)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 5. 데이터 출처
    # ══════════════════════════════════════════════════════════
    _heading(doc, "데이터 출처 (Sources)", level=1)

    if source_urls:
        _heading(doc, "참조 URL", level=2, color=C_GRAY)
        for url in source_urls[:30]:
            p_url = doc.add_paragraph(style="List Bullet")
            r_url = p_url.add_run(url)
            r_url.font.size = Pt(9)
            r_url.font.color.rgb = C_CYAN

    if sources:
        _heading(doc, "수집 미디어", level=2, color=C_GRAY)
        for src in sorted(set(sources)):
            p_s = doc.add_paragraph(style="List Bullet")
            p_s.add_run(src).font.size = Pt(10)

    # 푸터
    doc.add_paragraph()
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot = p_foot.add_run(
        f"Generated by Multi-LLM Strategy Analyzer\n"
        f"AI 기여: {attribution}  |  {now_utc}"
    )
    r_foot.font.size = Pt(8)
    r_foot.font.italic = True
    r_foot.font.color.rgb = C_GRAY

    # ── 저장 ──────────────────────────────────────────────────
    safe_date   = date_str.replace("-", "")
    filename    = f"semiconductor_brief_{report_type.lower()}_{safe_date}.docx"
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    filepath    = str(output_path / filename)

    doc.save(filepath)
    print(f"[Word] 저장 완료: {filepath}")
    return filepath
